from __future__ import annotations

import io
import re
from pathlib import Path


def build_report_context(report_type: str, project_info: dict, top_k: int = 8) -> str:
    query = f"{report_type} {project_info.get('description', '')} {project_info.get('project_name', '')}"
    from tools.file_tools import get_reference_context
    return get_reference_context(query, n_results=top_k)


def render_skeleton(report_type: str, project_info: dict) -> str:
    from core.template_engine import TemplateEngine
    engine = TemplateEngine()
    return engine.render_skeleton(report_type, project_info)


def count_tokens(text: str) -> int:
    try:
        import tiktoken
        enc = tiktoken.get_encoding("cl100k_base")
        return len(enc.encode(text))
    except Exception:
        return len(text) // 4


def trim_to_token_budget(text: str, budget: int) -> str:
    import config
    if count_tokens(text) <= budget:
        return text
    ratio = budget / max(count_tokens(text), 1)
    char_limit = int(len(text) * ratio * 0.95)
    return text[:char_limit] + "\n\n[Reference content trimmed to fit context window]"


def markdown_to_docx(markdown_text: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = markdown_text.split("\n")

    for line in lines:
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("---"):
            doc.add_paragraph("─" * 60)
        elif line.startswith("| "):
            _handle_table_line(doc, line, lines)
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            doc.add_paragraph(line.strip()[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", line.strip()):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", line.strip()), style="List Number")
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            cleaned = re.sub(r"\*(.+?)\*", r"\1", cleaned)
            cleaned = re.sub(r"`(.+?)`", r"\1", cleaned)
            doc.add_paragraph(cleaned)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _handle_table_line(doc, line: str, all_lines: list[str]):
    if re.match(r"^\|[\s\-|]+\|$", line.strip()):
        return
    cells = [c.strip() for c in line.strip().strip("|").split("|")]
    if not any(cells):
        return
    try:
        table = doc.tables[-1]
        row = table.add_row()
        for i, cell_text in enumerate(cells):
            if i < len(row.cells):
                row.cells[i].text = cell_text
    except (IndexError, AttributeError):
        table = doc.add_table(rows=1, cols=len(cells))
        table.style = "Table Grid"
        for i, cell_text in enumerate(cells):
            table.rows[0].cells[i].text = cell_text